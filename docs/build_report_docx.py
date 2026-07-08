"""Generate the FloodCast academic report as a Word (.docx) document.

Run:  python docs/build_report_docx.py
Produces: docs/FloodCast_Report.docx
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

OUT = Path(__file__).resolve().parent / "FloodCast_Report.docx"

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
MUTED = RGBColor(0x55, 0x55, 0x55)


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    return h


def para(doc, text, *, italic=False, size=None, color=None, align=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    return p


def numbered(doc, text):
    p = doc.add_paragraph(text, style="List Number")
    p.paragraph_format.space_after = Pt(2)
    return p


def equation(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.name = "Cambria Math"
    run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    return p


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        r = hdr[i].paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(val))
            r.font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def build():
    doc = Document()

    # Base style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # ---- Title ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("A Hybrid Physics–Machine-Learning System for Short-Lead "
                      "Urban Flood Nowcasting")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = ACCENT
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Methodology, Results and Discussion — Case study: South C, "
                     "Nairobi (rain gauge TA00026)")
    rs.italic = True
    rs.font.size = Pt(12)
    rs.font.color.rgb = MUTED
    doc.add_paragraph()

    # ---- Abstract ----
    add_heading(doc, "Abstract", 1)
    para(doc,
         "Urban pluvial (rainfall-driven) flooding in rapidly growing African "
         "cities is frequent, highly localised and poorly served by conventional "
         "hydrological warning systems. This report presents FloodCast, a "
         "prototype hybrid forecasting system that combines a lumped "
         "physically-based runoff–capacity model (the Generalized Runoff "
         "Capacity, GRC, surrogate) with an uncertainty-aware machine-learning "
         "(ML) residual learner to produce probabilistic 1–6 hour flood nowcasts "
         "for the South C area of Nairobi. Because no observed inundation record "
         "exists for the study area, model skill is demonstrated against a "
         "physically-based synthetic reference driven by the real observed "
         "rainfall. On a chronologically held-out test set, the hybrid model "
         "reduces the 1-hour flood-depth RMSE from 2.19 mm (physics only) to "
         "0.73 mm — a 67% reduction — achieves a flood-event discrimination score "
         "of AUC = 0.871 and a well-calibrated 90% predictive interval (empirical "
         "coverage 0.888). A cost–loss decision analysis shows the forecast "
         "delivers positive economic value over climatology at short lead times "
         "(best relative economic value 0.582 at 1 h). The results establish the "
         "method — not validated operational skill — and provide a template for "
         "deployment once real inundation observations become available.")

    # ---- 1. Introduction ----
    add_heading(doc, "1. Introduction", 1)
    add_heading(doc, "1.1 Problem context", 2)
    para(doc,
         "Nairobi's South C neighbourhood experiences recurrent flash flooding "
         "during the March–May long rains and October–December short rains. "
         "Flooding is driven by short, intense convective rainfall bursts "
         "interacting with undersized and frequently blocked drainage "
         "infrastructure, producing localised ponding at a small number of "
         "hydraulic chokepoints (culverts, outfalls, low-lying junctions). Two "
         "properties make this problem difficult for classical methods:")
    numbered(doc, "Rarity and class imbalance. Damaging flooding occupies a tiny "
                  "fraction of the record (here a base rate of 0.32% of hours), so "
                  "naïve models are biased toward \u201cno flood\u201d.")
    numbered(doc, "Spatial heterogeneity. Flood response varies sharply between "
                  "chokepoints, which a single lumped conceptual model cannot "
                  "represent.")

    add_heading(doc, "1.2 Rationale for a hybrid approach", 2)
    para(doc,
         "Purely physical (conceptual/hydrodynamic) models encode mass balance "
         "and capacity constraints but are structurally too simple (when lumped) "
         "or too data-hungry and slow (when fully distributed) for real-time "
         "neighbourhood nowcasting. Purely data-driven models can capture complex "
         "nonlinear patterns but ignore physical constraints, extrapolate poorly "
         "and are opaque. We therefore adopt a residual (delta) learning design: "
         "a fast physical surrogate provides a physically-consistent first guess, "
         "and an ML model learns the systematic residual between that surrogate "
         "and reality, while also quantifying predictive uncertainty. This "
         "preserves physical plausibility, concentrates the learning problem on "
         "what the physics misses, and yields calibrated probabilities suitable "
         "for risk-based decision-making.")

    add_heading(doc, "1.3 Objectives", 2)
    bullet(doc, "Develop a reproducible pipeline that forecasts flood depth and "
                "flood probability at lead times of 1–6 hours.")
    bullet(doc, "Quantify predictive uncertainty and decompose it into "
                "interpretable components (aleatory vs epistemic).")
    bullet(doc, "Evaluate not only statistical skill but decision value through a "
                "cost–loss framework.")
    bullet(doc, "Deliver the results through an operational-style dashboard "
                "(alerts, spatial risk map, printable bulletin).")

    # ---- 2. Data ----
    add_heading(doc, "2. Data", 1)
    table(doc,
          ["Source", "Variable", "Resolution", "Role"],
          [["Gauge TA00026", "Precipitation", "Hourly", "Primary model forcing"],
           ["CHIRPS satellite", "Precipitation", "Daily", "Regional-wetness context"],
           ["MAF metadata", "—", "—", "Station reference"]])
    para(doc,
         "The cleaned hourly gauge series spans 88,080 hours (~10.1 years), with "
         "8,368.5 mm of total rainfall over 12,191 wet hours and a peak intensity "
         "of 112.3 mm/h. Data preparation enforces a continuous, gap-free hourly "
         "index (missing hours set to zero rainfall and clipped to be "
         "non-negative), and repairs corrupted day-of-year–style dates in the "
         "CHIRPS spreadsheet (e.g. 2026-05-149) by reconstructing the calendar "
         "date from the stated year and day-of-year. CHIRPS daily totals are "
         "broadcast to each hour of the corresponding day and used only as a "
         "low-frequency wetness context feature, never as the sub-hourly signal.")

    # ---- 3. Methodology ----
    add_heading(doc, "3. Methodology", 1)
    para(doc,
         "The pipeline has five stages: (i) a synthetic physically-based "
         "reference; (ii) a GRC physical surrogate; (iii) feature engineering; "
         "(iv) an uncertainty-aware ML residual learner; and (v) ensemble "
         "forecasting with decision analysis.")
    para(doc,
         "Pipeline: Observed hourly rainfall → {Synthetic reference (6 chokepoint "
         "reservoirs); GRC surrogate (lumped reservoir); Feature engineering}. "
         "Residual targets = (truth − GRC) at leads 1–6 h → ML residual ensemble "
         "(+ variance heads) → physical projection (flood = GRC + residual) → "
         "ensemble & uncertainty decomposition → decision analysis (alerts, "
         "cost–loss, spatial map).",
         italic=True, color=MUTED)

    add_heading(doc, "3.1 Synthetic reference (\u201creality\u201d)", 2)
    para(doc,
         "In the absence of observed inundation data, a heterogeneous synthetic "
         "truth is generated from the real rainfall. The catchment is represented "
         "as six sub-catchment storage–overflow reservoirs (\u201cchokepoints\u201d), "
         "each with independently sampled parameters: runoff coefficient, "
         "infiltration loss, conveyance (pipe) capacity, storage capacity and a "
         "travel-time lag. For each sub-catchment j the effective runoff includes "
         "a mild nonlinear intensity boost,")
    equation(doc, "e(j,t) = max( c_r,j · R_lag(j,t) · (1 + 0.02·R_lag(j,t)) − f_j , 0 )")
    para(doc,
         "and the reservoir evolves with capacity-limited conveyance, with "
         "surface flooding equal to the overflow above storage capacity S_cap,j:")
    equation(doc, "S(j,t) = S(j,t−1) + e(j,t) − min(C_j, S(j,t−1)+e(j,t));   "
                  "d(j,t) = max(0, S(j,t) − S_cap,j)")
    para(doc,
         "The hourly flood signal emphasises the worst-affected locations by "
         "averaging the two deepest chokepoints per hour, and a heteroskedastic "
         "observation noise (variance increasing with flood magnitude) is added:")
    equation(doc, "y(t) = max(0, mean_top2(d(·,t)) + ε(t));   "
                  "ε(t) ~ N(0, σ0·(0.3 + flood(t)))")
    para(doc,
         "A binary flood event is defined by a depth threshold of 2.0 mm. This "
         "richer, noisy, spatially heterogeneous process is deliberately more "
         "complex than the lumped surrogate can represent, leaving a structured "
         "residual for the ML learner to recover. All reported skill is relative "
         "to this synthetic reference and therefore demonstrates method behaviour, "
         "not validated real-world performance.")

    add_heading(doc, "3.2 GRC physical surrogate", 2)
    para(doc,
         "The Generalized Runoff Capacity model is a single lumped reservoir "
         "integrating a physically-plausible mass balance with an infiltration "
         "loss and a routing lag:")
    equation(doc, "S(t) = S(t−1) + max(c_r·R_lag(t) − f, 0) − Q(t);   "
                  "Q(t) = min(min(C_g, C_p), S(t));   flood(t) = max(0, S(t) − S_cap)")
    para(doc,
         "By construction the per-step mass-balance residual is ~0 (verified "
         "numerically), and the model exposes internal states (storage, "
         "discharge, flood) that serve as physics priors for the ML stage. Being "
         "lumped, it cannot capture the multi-chokepoint heterogeneity of the "
         "reference — the intended source of the learnable residual.")

    add_heading(doc, "3.3 Feature engineering", 2)
    para(doc, "Features (~20 predictors) combine rainfall history and hydrologic "
              "state:")
    bullet(doc, "Rolling accumulations over windows of 1, 2, 3, 6, 12, 24 h.")
    bullet(doc, "Burst intensity (rolling maxima over 3, 6, 12 h).")
    bullet(doc, "Antecedent Precipitation Index (exponential wetness memory, "
                "decay k = 0.9).")
    bullet(doc, "Dry-spell length (hours since last rain, capped at 168 h).")
    bullet(doc, "CHIRPS regional-wetness context.")
    bullet(doc, "GRC internal state (flood, storage, discharge) — the physics prior.")
    bullet(doc, "Seasonality/diurnality harmonics (sin/cos of hour-of-day and "
                "day-of-year).")
    para(doc,
         "The prediction target at lead h is the residual between the reference "
         "and the GRC flood, shifted forward: r(t+h) = y(t+h) − flood_GRC(t+h), "
         "for h in {1,…,6} h.")

    add_heading(doc, "3.4 Uncertainty-aware ML residual learner", 2)
    para(doc,
         "The residual is learned by a deep-ensemble style estimator, one per "
         "lead time. It comprises:")
    bullet(doc, "M = 8 gradient-boosted members (HistGradientBoostingRegressor; "
                "max_depth = 6, max_iter = 250, learning_rate = 0.06, "
                "l2_regularization = 1.0), each trained on an independent bootstrap "
                "resample with a distinct random seed. The spread of member means "
                "provides the epistemic (model/parameter) variance.")
    bullet(doc, "An aleatory-variance head: a separate regressor trained on the "
                "log of the squared ensemble-mean error, giving a heteroskedastic "
                "(input-dependent) data-noise variance.")
    para(doc, "Predictive moments follow the law of total variance:")
    equation(doc, "μ = (1/M) Σ_m μ_m;   "
                  "σ²_total = (1/M) Σ_m σ²_m  [aleatory]  +  Var_m(μ_m)  [epistemic]")
    para(doc,
         "The residual mean is added back to the GRC flood and projected onto the "
         "feasible set (non-negativity, a soft realisation of the physical "
         "mass-balance/capacity constraints): ŷ = max(0, flood_GRC + μ). The "
         "flood-exceedance probability is obtained analytically from the Gaussian "
         "predictive distribution, P(ŷ > τ) = 1 − Φ((τ − μ)/σ).")

    add_heading(doc, "3.5 Ensemble forecasting and uncertainty decomposition", 2)
    para(doc,
         "For a single forecast instance, a full Monte-Carlo ensemble propagates "
         "three uncertainty sources:")
    numbered(doc, "Forcing uncertainty — 25 multiplicative lognormal rainfall "
                  "perturbations (coefficient of variation 0.35);")
    numbered(doc, "Parameter/epistemic uncertainty — 25 GRC parameter draws "
                  "(runoff coefficient, pipe capacity, storage capacity);")
    numbered(doc, "Model epistemic uncertainty — the 8 ML ensemble members;")
    para(doc,
         "yielding 25 × 25 × 8 = 5,000 predictive draws, whose histogram and "
         "exceedance probability are surfaced in the dashboard.")

    add_heading(doc, "3.6 Evaluation protocol", 2)
    para(doc,
         "Models are trained on the first 80% of the record and evaluated on the "
         "final 20% (a strict chronological hold-out to avoid look-ahead leakage; "
         "seed fixed at 42 for reproducibility). Skill is assessed with:")
    bullet(doc, "Deterministic: MAE, RMSE, bias.")
    bullet(doc, "Probabilistic: CRPS (closed-form Gaussian), 90% interval "
                "coverage, Brier score, ROC-AUC for event discrimination.")
    bullet(doc, "Decision value: the relative economic value (REV) from a "
                "cost–loss model. With loss normalised to L = 1 and cost C, the "
                "optimal action rule is \u201cact when p > p* = C/L\u201d, and")
    equation(doc, "REV = (E_clim − E_forecast) / (E_clim − E_perfect)")
    para(doc, "evaluated across cost–loss ratios {0.02, 0.05, 0.1, 0.2, 0.3, 0.5, "
              "0.7}.")

    # ---- 4. Results ----
    add_heading(doc, "4. Results", 1)
    add_heading(doc, "4.1 Headline skill (1-hour lead)", 2)
    table(doc,
          ["Metric", "Physics only (GRC)", "Hybrid (GRC + ML)", "Change"],
          [["Flood-depth RMSE (mm)", "2.188", "0.730", "−67%"],
           ["Event discrimination (AUC)", "—", "0.871", "—"],
           ["90% interval coverage", "—", "0.888", "≈ nominal"],
           ["Best relative economic value", "—", "0.582", "vs climatology"]])
    para(doc,
         "The residual learner removes roughly two-thirds of the physics-only "
         "error, indicating that a large, systematic (learnable) component of the "
         "reference response is missed by the lumped surrogate and successfully "
         "recovered by the ML stage. The near-nominal 90% coverage (0.888) shows "
         "the predictive intervals are well calibrated — neither over- nor "
         "under-confident.")

    add_heading(doc, "4.2 Skill as a function of lead time", 2)
    table(doc,
          ["Lead (h)", "RMSE (mm)", "CRPS", "AUC", "Brier", "Coverage90", "Best REV"],
          [["1", "0.730", "0.087", "0.871", "0.0030", "0.888", "0.582"],
           ["2", "0.746", "0.092", "0.767", "0.0029", "0.900", "0.443"],
           ["3", "0.818", "0.104", "0.711", "0.0044", "0.906", "0.138"],
           ["4", "0.840", "0.114", "0.652", "0.0054", "0.915", "−0.049"],
           ["5", "0.873", "0.119", "0.645", "0.0060", "0.910", "−0.072"],
           ["6", "0.927", "0.123", "0.644", "0.0066", "0.907", "−0.055"]])
    para(doc,
         "Skill degrades monotonically with horizon, as expected: RMSE and CRPS "
         "rise while AUC falls from 0.871 (1 h) to ~0.64 (6 h). Crucially, "
         "positive decision value is confined to the 1–3 h window (REV > 0), "
         "beyond which the forecast no longer beats climatology at any cost–loss "
         "ratio for this synthetic reference. Calibration (coverage ≈ 0.89–0.92) "
         "remains stable across all leads.")

    add_heading(doc, "4.3 Uncertainty decomposition", 2)
    table(doc,
          ["Lead (h)", "Aleatory (mm²)", "Epistemic (mm²)", "Epistemic share"],
          [["1", "0.055", "0.080", "59%"],
           ["2", "0.073", "0.168", "70%"],
           ["3", "0.040", "0.048", "55%"],
           ["4", "0.023", "0.187", "89%"],
           ["5", "0.040", "0.228", "85%"],
           ["6", "0.048", "0.281", "85%"]])
    para(doc,
         "Epistemic (model/parameter) uncertainty dominates and grows with lead "
         "time, while aleatory (irreducible data-noise) uncertainty stays "
         "comparatively flat. This is diagnostically useful: it implies that "
         "longer-lead skill is limited mainly by reducible model uncertainty and "
         "could be improved with more/better training data, richer features or a "
         "larger ensemble — rather than by an irreducible noise floor.")

    add_heading(doc, "4.4 Operational outputs", 2)
    para(doc,
         "The dashboard translates the statistical forecast into operational "
         "products: a four-tier alert scheme (All-clear / Watch / Warning / "
         "Emergency) keyed to exceedance probability; a per-chokepoint spatial "
         "risk map over the six South C nodes (weighted by climatological "
         "susceptibility); a lead-time risk profile; and an auto-generated, "
         "printable alert bulletin summarising status, most-exposed chokepoints, "
         "recommended actions and hold-out verification statistics.")

    # ---- 5. Discussion ----
    add_heading(doc, "5. Discussion", 1)
    add_heading(doc, "5.1 Interpretation", 2)
    para(doc,
         "The central finding is that residual hybridisation is highly effective: "
         "a deliberately simple physical model, corrected by an uncertainty-aware "
         "ML learner, reduces error by two-thirds and produces calibrated "
         "probabilities. The physics supplies structure, feasibility and "
         "interpretable state variables (which also serve as informative "
         "features), while the ML captures the nonlinear, heterogeneous "
         "chokepoint behaviour the lumped model omits. Retaining the GRC prior "
         "and projecting predictions onto physically feasible values guards "
         "against the unphysical extrapolation typical of unconstrained ML.")
    para(doc,
         "The uncertainty decomposition adds practical value beyond a single "
         "accuracy number: separating aleatory from epistemic variance tells an "
         "operator why confidence is low at a given lead and points to the most "
         "effective route to improvement (here, reducing epistemic uncertainty at "
         "longer leads).")
    para(doc,
         "The decision-oriented evaluation is arguably the most policy-relevant "
         "result. Statistical skill and usefulness are not the same thing under "
         "severe class imbalance; the cost–loss/REV analysis shows precisely where "
         "the forecast is worth acting on (1–3 h) and where it is not (≥ 4 h). "
         "This directly informs how far ahead warnings should be issued.")

    add_heading(doc, "5.2 Comparison with conventional approaches", 2)
    para(doc,
         "Relative to a lumped conceptual model alone, the hybrid adds accuracy, "
         "calibrated uncertainty and event-discrimination skill. Relative to a "
         "black-box ML model, it adds physical consistency, interpretable "
         "intermediate states and robustness to extrapolation. The architecture "
         "is also computationally light: once trained, per-hour forecasts are "
         "closed-form, and the pretrained model state is serialised so the "
         "deployed service starts in seconds without retraining.")

    add_heading(doc, "5.3 Limitations", 2)
    numbered(doc, "Synthetic ground truth. The most important caveat: skill is "
                  "measured against a physically-based synthetic reference, not "
                  "observed inundation. The numbers validate the methodology and "
                  "calibration, not operational accuracy.")
    numbered(doc, "Single gauge forcing. One rain gauge cannot resolve the strong "
                  "spatial variability of convective storms; the spatial map is an "
                  "informed disaggregation, not an independently observed field.")
    numbered(doc, "Idealised chokepoints. Locations and susceptibilities are "
                  "representative and for demonstration; real drainage geometry, "
                  "blockages and backwater effects are not explicitly modelled.")
    numbered(doc, "Stationarity assumptions. Chronological hold-out mitigates "
                  "leakage but the model assumes the rainfall–flood relationship is "
                  "stationary, which urbanisation and climate change may violate.")
    numbered(doc, "Short useful horizon. Positive decision value currently extends "
                  "only to ~3 h, limiting lead time for some response actions.")

    add_heading(doc, "5.4 Future work", 2)
    bullet(doc, "Real observations: integrate crowd-sourced/CCTV/IoT depth reports "
                "or SWMM hydrodynamic output to replace the synthetic target and "
                "enable genuine validation.")
    bullet(doc, "Spatial forcing: assimilate radar or satellite QPE and multiple "
                "gauges for a truly distributed nowcast.")
    bullet(doc, "Extended lead time: couple with numerical weather prediction / "
                "rainfall nowcasting to push useful horizons beyond 3 h.")
    bullet(doc, "Richer learners: temporal deep models (e.g. sequence models) and "
                "probabilistic calibration refinements.")
    bullet(doc, "Field trial: pilot the alerting workflow with the responsible "
                "drainage/disaster-management authority and measure realised "
                "decision value.")

    # ---- 6. Conclusion ----
    add_heading(doc, "6. Conclusion", 1)
    para(doc,
         "This work demonstrates a reproducible, uncertainty-aware hybrid "
         "physics–ML nowcasting system for short-lead urban pluvial flooding in "
         "South C, Nairobi. By learning the residual of a lightweight physical "
         "runoff–capacity surrogate, the system cuts 1-hour flood-depth RMSE by "
         "67% (2.19 → 0.73 mm), discriminates flood events well (AUC 0.871), "
         "produces well-calibrated 90% predictive intervals (coverage 0.888), and "
         "— through a cost–loss analysis — delivers positive economic value over "
         "climatology at 1–3 h lead (best REV 0.582). The explicit "
         "aleatory/epistemic uncertainty decomposition and the decision-focused "
         "evaluation make the outputs both interpretable and operationally "
         "actionable via an alerting dashboard, spatial risk map and printable "
         "bulletin. While the present skill is established against a synthetic "
         "reference and must be revalidated on real inundation data, the "
         "methodology provides a practical, physically-grounded and "
         "computationally efficient blueprint for operational urban flood "
         "early-warning in data-scarce cities.")

    add_heading(doc, "Reproducibility", 2)
    para(doc,
         "The full pipeline, pretrained model artifact and dashboard are available "
         "in the project repository. All results in this report were produced with "
         "the committed model state (random seed 42, 80/20 chronological split). "
         "To regenerate the model artifact after any model or data change, run "
         "\u201cpython -m app.service\u201d from the prototype folder and commit the "
         "updated file.")
    p = doc.add_paragraph()
    rr = p.add_run("Disclaimer. Flood targets are SWMM-style synthetic values. "
                   "Results demonstrate method and calibration, not validated "
                   "real-world skill, and the system is a research prototype — not "
                   "an official flood warning.")
    rr.italic = True
    rr.font.size = Pt(9.5)
    rr.font.color.rgb = MUTED

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
